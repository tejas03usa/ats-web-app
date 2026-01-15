import streamlit as st
import os
import json
import io
import re
import urllib.parse
from openai import OpenAI
from xml.sax.saxutils import escape as xml_escape
from dotenv import load_dotenv

# Load .env variables (if running locally with .env)
load_dotenv()

# File Processing
import PyPDF2 as pdf
import docx
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# --- Page Configuration ---
st.set_page_config(
    page_title="Career Toolkit",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Session State Initialization ---
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'cover_letter' not in st.session_state:
    st.session_state.cover_letter = ""
if 'study_plan' not in st.session_state:
    st.session_state.study_plan = ""
if 'qa_result' not in st.session_state:
    st.session_state.qa_result = ""
if 'search_config' not in st.session_state:
    st.session_state.search_config = None

# --- Helper Functions ---

def check_password():
    """Checks the password against Streamlit Secrets or defaults for local dev."""
    def password_entered():
        # Check if password matches strict secret or default dev password
        # Use try/except to handle missing secrets.toml locally
        try:
            expected_password = st.secrets.get("APP_PASSWORD", "admin123")
        except Exception:
            expected_password = "admin123" # Fallback for local dev if no secrets
            
        if st.session_state["password_input"] == expected_password:
            st.session_state.authenticated = True
            del st.session_state["password_input"]
        else:
            st.session_state.authenticated = False
            st.error("😕 Incorrect Password")

    if not st.session_state.authenticated:
        st.text_input("Enter App Password to Access", type="password", on_change=password_entered, key="password_input")
        return False
    return True

def extract_text_from_pdf(uploaded_file):
    text = ""
    reader = pdf.PdfReader(uploaded_file)
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    return '\n'.join([para.text for para in doc.paragraphs])

def clean_resume_text(text):
    """Removes 'Resume', 'CV' headers from the top."""
    return re.sub(r'^(#+\s*)?(Resume|CV|Curriculum Vitae|Optimized Resume)\:?\s*\n+', '', text, flags=re.IGNORECASE).strip()

# --- Export Functions (Adapted for Streamlit Memory Buffers) ---

def create_docx(content, title_prefix):
    doc = Document()
    if "Resume" not in title_prefix:
        doc.add_heading(title_prefix, 0)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('###') or line.startswith('##'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold logic
            parts = re.split(r'(\*\*.*?\*\*)', line[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(content, title_prefix):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ModernHeader', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, spaceAfter=6, spaceBefore=12, textColor=colors.darkslategray))
    styles.add(ParagraphStyle(name='ModernNormal', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14))
    styles.add(ParagraphStyle(name='ModernBullet', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, leftIndent=20))

    story = []
    if "Resume" not in title_prefix:
        story.append(Paragraph(title_prefix, styles["Heading1"]))
        story.append(Spacer(1, 12))

    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        line_safe = xml_escape(line)
        
        if line_safe.startswith('###') or line_safe.startswith('##'):
            clean_text = line_safe.replace('#', '').strip()
            story.append(Paragraph(clean_text, styles["ModernHeader"]))
        elif line_safe.startswith('- ') or line_safe.startswith('* '):
            clean_text = line_safe[2:].strip()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_text)
            story.append(Paragraph(f"• {clean_text}", styles["ModernBullet"]))
        else:
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line_safe)
            story.append(Paragraph(clean_text, styles["ModernNormal"]))
        story.append(Spacer(1, 3))

    doc.build(story)
    buffer.seek(0)
    return buffer

# --- Main App ---

def main():
    if not check_password():
        st.stop() # Stop execution if not authenticated

    # Sidebar Configuration
    with st.sidebar:
        st.title("🧰 Toolkit Setup")
        
        # API Key Handling
        user_input_key = st.text_input("OpenAI API Key", type="password", help="Enter your sk-... key")
        
        # Logic to determine effective API Key
        # Priority: 1. User Input -> 2. Streamlit Secrets -> 3. Local .env
        api_key = user_input_key
        
        if not api_key:
            try:
                api_key = st.secrets.get("OPENAI_API_KEY", "")
            except Exception:
                pass # Secrets not found
        
        if not api_key:
            api_key = os.getenv("OPENAI_API_KEY", "")

        # UI Feedback for Key Status
        if api_key and not user_input_key:
            st.success("✅ Loaded API Key from Secrets/Env")
        elif not api_key:
            st.warning("⚠️ No API Key found. Please enter one.")
        
        st.divider()
        
        st.subheader("1. Upload Data")
        uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.pdf'):
                    st.session_state.resume_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.name.endswith('.docx'):
                    st.session_state.resume_text = extract_text_from_docx(uploaded_file)
                st.success("Resume Loaded!")
            except Exception as e:
                st.error(f"Error: {e}")

        jd_text = st.text_area("Job Description", height=150, placeholder="Paste JD here...")
        location = st.text_input("Target Location", placeholder="e.g. New York, NY or Remote")

    # Main Area
    st.title("🚀 Smart Career Toolkit")
    
    if not api_key:
        st.info("👈 Please enter your OpenAI API Key in the sidebar to unlock the tools.")
        st.stop()

    # Tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📊 ATS Scan", 
        "✨ Optimized Resume", 
        "✉️ Cover Letter", 
        "🎓 Interview Prep", 
        "❓ Q&A", 
        "🔍 Job Search"
    ])

    # --- TAB 1: ATS SCAN ---
    with tab1:
        if st.button("Run ATS Analysis", type="primary"):
            if not st.session_state.resume_text or not jd_text:
                st.error("Please upload a resume and paste a Job Description.")
            else:
                with st.spinner("Analyzing resume..."):
                    try:
                        client = OpenAI(api_key=api_key)
                        prompt = f"""
                        You are an expert ATS optimizer.
                        RESUME: {st.session_state.resume_text[:8000]}
                        JD: {jd_text[:4000]}
                        Task: 
                        1. Analyze resume vs JD. 
                        2. Rewrite resume to include missing keywords naturally (100% match goal).
                        3. Calculate match score (0-100).
                        4. List missing keywords.
                        
                        Return JSON: {{ "analysis": "string", "optimized_resume": "string", "score": int, "missing_keywords": ["kw1"] }}
                        IMPORTANT: "optimized_resume" must contain ONLY the resume content. Start with Name. No "Resume" headers.
                        """
                        response = client.chat.completions.create(
                            model="gpt-4o", messages=[{"role": "user", "content": prompt}], response_format={"type": "json_object"}
                        )
                        st.session_state.analysis_result = json.loads(response.choices[0].message.content)
                        st.session_state.analysis_result['optimized_resume'] = clean_resume_text(st.session_state.analysis_result['optimized_resume'])
                    except Exception as e:
                        st.error(f"Error: {e}")

        if st.session_state.analysis_result:
            res = st.session_state.analysis_result
            col1, col2 = st.columns([1, 3])
            with col1:
                st.metric("Match Score", f"{res.get('score', 0)}%")
            with col2:
                st.write("**Missing Keywords:**")
                st.write(", ".join(res.get('missing_keywords', [])))
            st.markdown("### Detailed Analysis")
            st.write(res.get('analysis', ''))

    # --- TAB 2: OPTIMIZED RESUME ---
    with tab2:
        if st.session_state.analysis_result:
            opt_text = st.session_state.analysis_result.get('optimized_resume', '')
            st.text_area("Preview", opt_text, height=500)
            
            c1, c2 = st.columns(2)
            with c1:
                st.download_button("Download PDF", data=create_pdf(opt_text, "Optimized Resume"), file_name="Optimized_Resume.pdf", mime="application/pdf")
            with c2:
                st.download_button("Download Word", data=create_docx(opt_text, "Optimized Resume"), file_name="Optimized_Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("Run ATS Analysis first.")

    # --- TAB 3: COVER LETTER ---
    with tab3:
        if st.button("Generate Cover Letter"):
            if not st.session_state.resume_text or not jd_text:
                st.error("Missing inputs.")
            else:
                with st.spinner("Drafting letter..."):
                    client = OpenAI(api_key=api_key)
                    prompt = f"""Write a professional cover letter based on this RESUME and JD. Professional tone. Concise.
                    RESUME: {st.session_state.resume_text[:6000]}
                    JD: {jd_text[:3000]}"""
                    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
                    st.session_state.cover_letter = resp.choices[0].message.content
        
        if st.session_state.cover_letter:
            st.text_area("Preview", st.session_state.cover_letter, height=400)
            st.download_button("Download PDF", data=create_pdf(st.session_state.cover_letter, "Cover Letter"), file_name="Cover_Letter.pdf", mime="application/pdf")

    # --- TAB 4: INTERVIEW PREP ---
    with tab4:
        if st.button("Generate Study Plan"):
            if not jd_text: st.error("Please provide a Job Description.")
            else:
                with st.spinner("Creating plan..."):
                    client = OpenAI(api_key=api_key)
                    prompt = f"Create a Markdown Interview Study Plan for this JD: {jd_text[:4000]}. Include topics and YouTube search terms."
                    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
                    st.session_state.study_plan = resp.choices[0].message.content
        
        if st.session_state.study_plan:
            st.markdown(st.session_state.study_plan)
            st.download_button("Download PDF", data=create_pdf(st.session_state.study_plan, "Study Plan"), file_name="Study_Plan.pdf", mime="application/pdf")

    # --- TAB 5: Q&A ---
    with tab5:
        if st.button("Generate Interview Q&A"):
            if not st.session_state.resume_text: st.error("Please upload Resume.")
            else:
                with st.spinner("Generating questions..."):
                    client = OpenAI(api_key=api_key)
                    prompt = f"Generate 20 interview questions and answers based on this Resume and JD. Format in Markdown. RESUME: {st.session_state.resume_text[:6000]} JD: {jd_text[:3000]}"
                    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
                    st.session_state.qa_result = resp.choices[0].message.content
        
        if st.session_state.qa_result:
            st.markdown(st.session_state.qa_result)
            st.download_button("Download PDF", data=create_pdf(st.session_state.qa_result, "Interview Q&A"), file_name="Interview_QA.pdf", mime="application/pdf")

    # --- TAB 6: JOB SEARCH ---
    with tab6:
        if st.button("Configure Smart Search"):
            if not st.session_state.resume_text or not location:
                st.error("Resume and Location required.")
            else:
                with st.spinner("Optimizing search terms..."):
                    client = OpenAI(api_key=api_key)
                    prompt = f"""
                    Analyze resume for job search.
                    RESUME: {st.session_state.resume_text[:2000]}
                    LOCATION: {location}
                    Return JSON: {{ "job_titles": ["title1", "title2"], "keywords": ["kw1", "kw2"], "location": "string" }}
                    """
                    resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}], response_format={"type": "json_object"})
                    st.session_state.search_config = json.loads(resp.choices[0].message.content)

        if st.session_state.search_config:
            cfg = st.session_state.search_config
            titles_query = f"({' OR '.join(cfg['job_titles'])})"
            keywords_query = f"({' OR '.join(cfg['keywords'])})"
            loc = cfg['location']
            
            st.success(f"Optimized for: **{', '.join(cfg['job_titles'])}** in **{loc}**")
            
            # Helper to create links
            def make_link(url, text):
                st.markdown(f'<a href="{url}" target="_blank" style="text-decoration:none;"><button style="background-color:#0077b5;color:white;border:none;padding:10px 20px;border-radius:5px;cursor:pointer;width:100%;margin:5px 0;">{text} ↗</button></a>', unsafe_allow_html=True)

            # Encodings
            full_q = f"{titles_query} {keywords_query}"
            q_enc = urllib.parse.quote(full_q)
            l_enc = urllib.parse.quote(loc)
            
            # Google Specific
            g_q = full_q.replace('(', '').replace(')', '').replace(' OR ', ' ')
            g_q_enc = urllib.parse.quote_plus(g_q)
            g_l_enc = urllib.parse.quote_plus(loc)

            col_a, col_b = st.columns(2)
            with col_a:
                make_link(f"https://www.linkedin.com/jobs/search/?keywords={q_enc}&location={l_enc}", "Search LinkedIn")
                make_link(f"https://www.indeed.com/jobs?q={q_enc}&l={l_enc}", "Search Indeed")
            with col_b:
                make_link(f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={q_enc}&locKeyword={l_enc}", "Search Glassdoor")
                make_link(f"https://www.google.com/search?q={g_q_enc}+jobs+in+{g_l_enc}", "Search Google Jobs")

            st.text_input("Copy Search String", value=full_q)

if __name__ == "__main__":
    main()
